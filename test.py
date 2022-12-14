from types import NoneType
import requests
from bs4 import BeautifulSoup
from PIL import Image
import os.path
from urlhelpers import cut_link, post_from_link
from fake_useragent import UserAgent
from dotenv import load_dotenv
import os
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver import Keys
from docx import Document


options = Options()
options.add_argument('--headless')
options.add_argument('--disable-gpu')
options.add_argument('--no-sandbox')
        
service = Service(executable_path='D:\chromedriver.exe')
driver = webdriver.Chrome(service=service, options=options)   
    
urls_list = ["https://navalny.livejournal.com/2008/02/12/"] 
url = "https://navalny.livejournal.com/209488.html"
    
    
    
def save_post(url):
    count = 0
    dir = f"posts/{post_from_link(url)}"
    if not os.path.exists(dir):
        os.makedirs(dir)
    req = requests.get(url)
    soup = BeautifulSoup(req.text, "html.parser")
    heading = soup.find('h1')
    text = soup.find('div', class_="b-singlepost-bodywrapper")
    
    
    aricle = soup.select("div.b-singlepost-bodywrapper")
        
    for i in aricle:
        for link in i.find_all('img'):
            count += 1
            res = link.get('src')
            print(res)
            if '.png' in res:
                img = Image.open(requests.get(res, stream = True).raw)
                img.save(f"posts/{post_from_link(url)}/{count+1}.png")
            if '.jpg' in res:
                img = Image.open(requests.get(res, stream = True).raw)
                img.save(f"posts/{post_from_link(url)}/{count+1}.jpg")
    
    document = Document()
    try:
        document.add_heading(heading.get_text(), level=1)
    except AttributeError or TypeError:
        pass
    try:
        document.add_paragraph(text.get_text())
    except AttributeError or TypeError:
        pass

    document.save(f"posts/{post_from_link(url)}/post.docx")



save_post(url)