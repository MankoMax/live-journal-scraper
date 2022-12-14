import requests
from bs4 import BeautifulSoup
from PIL import Image
import os.path
from urlhelpers import cut_link, post_from_link
from fake_useragent import UserAgent
from dotenv import load_dotenv
import os
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver import Keys
from docx import Document

load_dotenv()

HTTP_PROXY = os.getenv("HTTP_PROXY")
HTTPS_PROXY = os.getenv("HTTPS_PROXY")

class Scrapper():
    
    
    def __init__(self, calendar_page):
        self.calendar_page = calendar_page
        self.urls = []
        self.ua = UserAgent()
        self.headers = {
            'User-Agent': self.ua.random
            }

        self.proxies = {
            'http': HTTP_PROXY,
            'https': HTTPS_PROXY,
            } 
        
        options = Options()
        options.add_argument('--headless')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
                
        service = Service(executable_path='D:\chromedriver.exe')
        self.driver = webdriver.Chrome(service=service, options=options)
                 
                        
    def get_calendar_years(self):
        urls = []
        req = requests.get(self.calendar_page, headers=self.headers, proxies=self.proxies)
        soup = BeautifulSoup(req.text, "html.parser")
        text = soup.select("div.entry-text")
        for i in text:
            for link in i.find_all('a'):
                res = cut_link(link.get('href'))
                if res not in urls:
                    urls.append(res)
        return urls


    def get_calendar_days(self, urls_list):
        urls = []
        for i in urls_list:
            self.driver.get(i)
            calendars = self.driver.find_element(By.CSS_SELECTOR, value="div.entry-text")
            bodies = calendars.find_elements(By.TAG_NAME, value="tbody")
            for body in bodies:
                days = body.find_elements(By.TAG_NAME, value="a")
                days = [day.get_attribute("href") for day in days]
                for day in days:
                    if day not in urls:
                        urls.append(day)
        return urls


    def get_posts_links(self, urls_list):
        urls = []
        for i in urls_list:
            self.driver.get(i)
            posts = self.driver.find_elements(By.CLASS_NAME, value="subj-link")
            posts = [post.get_attribute("href") for post in posts]
            for post in posts:
                if post not in urls:
                    urls.append(post)
        return urls


    def save_post(self, url):
        count = 0
        dir = f"posts/{post_from_link(url)}"
        if not os.path.exists(dir):
            os.makedirs(dir)
        req = requests.get(url, headers=self.headers, proxies=self.proxies)
        soup = BeautifulSoup(req.text, "html.parser")
        heading = soup.find('h1')
        text = soup.find('div', class_="b-singlepost-bodywrapper")
        aricle = soup.select("div.b-singlepost-bodywrapper")
            
        for i in aricle:
            for link in i.find_all('img'):
                count += 1
                res = link.get('src')
                if '.png' in res:
                    img = Image.open(requests.get(res, stream = True).raw)
                    img.save(f"posts/{post_from_link(url)}/{count+1}.png")
                if '.jpg' in res:
                    img = Image.open(requests.get(res, stream = True).raw)
                    img.save(f"posts/{post_from_link(url)}/{count+1}.jpg")
        
        document = Document()
        try:
            document.add_heading(f'{heading.get_text()}\n', level=1)
        except AttributeError or TypeError:
            pass
        try:
            document.add_paragraph(text.get_text(separator="\n") + f"\n\nСсылка на пост - {url}")
        except AttributeError or TypeError:
            pass

        document.save(f"posts/{post_from_link(url)}/post.docx")
        print(f"Пост сохранен - {url}")
        
        
        
        
        
        